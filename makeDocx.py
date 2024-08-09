from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx.enum.section import WD_ORIENT

import logging
from candidates import Candidate, list_candidates
from contests import Contest, list_contests
from datetime import datetime
import os


def candidates_in_contest(candidates, contest_id):
    # Filter candidates based on contest_id
    filtered_candidates = [candidate for candidate in candidates.values()
                           if candidate.contest_id == contest_id]

    # Sort the candidates by total_votes
    sorted_candidates = sorted(filtered_candidates, key=lambda candidate: candidate.total_votes, reverse=True)

    return sorted_candidates


def set_cell_font(cell, font_name, font_size, bold=False):
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_cell_background(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color)
    cell_properties.append(cell_shading)

def set_bottom_border(cell, border_color="000000", border_size="4"):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), border_size)
    bottom_border.set(qn('w:space'), '0')
    bottom_border.set(qn('w:color'), border_color)
    tc_borders.append(bottom_border)
    tc_pr.append(tc_borders)

def estimate_table_size(c_in_c, font_size=24, padding=2):
    num_rows = len(c_in_c) + 1  # Include header row
    estimated_row_height = font_size + padding
    total_height = num_rows * estimated_row_height
    return total_height


def generate_opendoc(contests, candidates, file_path):
    doc = Document()
    doc.add_heading('Election Results', 0)

    #estimate contest heights to insert page breaks
    # Set page layout (optional, if you need to switch to landscape mode)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section = doc.sections[-1]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    current_height = 0
    page_height_limit = 250  # wild estimate

    for contest_id, contest in contests.items():

        c_in_c = candidates_in_contest(candidates, contest_id)
        # print(f"contest {contest.datalink_value} len c_in_c: {len(c_in_c)} = {c_in_c}")
        estimated_height = estimate_table_size(c_in_c)

        # print(f"cur {current_height} + {estimated_height} = {current_height + estimated_height}>? {page_height_limit}")

        # Check if the contest fits on the current page
        if current_height + estimated_height > page_height_limit:
            doc.add_page_break()
            # print("adding break")
            current_height = 0  # Reset the current height for the new page
        # else:
        #     print("no need to add break")
        current_height += estimated_height  # Add the contest's height to the current page height
        # print(f"page height now {current_height}\n")

        heading = doc.add_heading(contest.datalink_value, level=1)

        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(18)  # Larger font size for headings
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            run_properties = run._r.get_or_add_rPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D3D3D3')
            run_properties.append(shading)

        table = doc.add_table(rows=1, cols=3)
        table.columns[0].width = Inches(4)  # Adjust as needed for long names
        table.columns[1].width = Inches(1)  # Narrow width for Votes
        table.columns[2].width = Inches(1)  # Narrow width for Percent

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Candidate'
        hdr_cells[1].text = 'Votes'
        hdr_cells[2].text = 'Percentage'

        set_cell_font(hdr_cells[0], 'Arial', 16, True)
        set_cell_font(hdr_cells[1], 'Arial', 16, True)
        set_cell_font(hdr_cells[2], 'Arial', 16, True)

        row_num = 1
        for candidate in c_in_c:
            row_cells = table.add_row().cells
            row_cells[0].text = candidate.candidate_name
            row_cells[1].text = str(candidate.total_votes)
            row_cells[2].text = f"{candidate.percent_votes:.1f}%"

            set_cell_font(row_cells[0], 'Arial', 16)
            set_cell_font(row_cells[1], 'Arial', 16)
            set_cell_font(row_cells[2], 'Arial', 16)

            # for cell in row_cells:
            #     set_bottom_border(cell)

            # if row_num % 2 == 1:
            #     for cell in row_cells:
            #         set_cell_background(cell, 'D3D3D3')  # Light grey color
            # row_num += 1

        #boi_paragraph = doc.add_paragraph(f"Blank, Over, Invalid Votes: {contest.bad_boi}")
        boi_percent_paragraph = doc.add_paragraph(f" percent total: {contest.percent_bad_boi}%")

        for paragraph in [boi_percent_paragraph]:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16)
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    doc.save(file_path)
    print(f"Saved Formated Results to {file_path}")
