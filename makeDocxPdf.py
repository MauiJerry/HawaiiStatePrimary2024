from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import win32print
import win32api


# Assuming the structure of the Candidate and Contest classes

class Candidate:
    def __init__(self, name, votes):
        self.name = name
        self.votes = votes


class Contest:
    def __init__(self, name, candidates, blank_votes, over_votes, invalid_votes):
        self.name = name
        self.candidates = candidates
        self.blank_votes = blank_votes
        self.over_votes = over_votes
        self.invalid_votes = invalid_votes


def set_cell_font(cell, font_name, font_size, bold=False):
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def generate_opendoc(contests):
    doc = Document()
    doc.add_heading('Election Results', 0)

    for contest in contests:
        heading = doc.add_heading(contest.name, level=1)
        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(18)  # Larger font size for headings
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Candidate'
        hdr_cells[1].text = 'Votes'
        hdr_cells[2].text = 'Percentage'

        set_cell_font(hdr_cells[0], 'Arial', 16, True)
        set_cell_font(hdr_cells[1], 'Arial', 16, True)
        set_cell_font(hdr_cells[2], 'Arial', 16, True)

        total_votes = sum(candidate.votes for candidate in contest.candidates)
        total_votes += contest.blank_votes + contest.over_votes + contest.invalid_votes

        sorted_candidates = sorted(contest.candidates, key=lambda x: x.votes, reverse=True)
        for candidate in sorted_candidates:
            row_cells = table.add_row().cells
            row_cells[0].text = candidate.name
            row_cells[1].text = str(candidate.votes)
            percentage = (candidate.votes / total_votes) * 100
            row_cells[2].text = f"{percentage:.2f}%"

            set_cell_font(row_cells[0], 'Arial', 16)
            set_cell_font(row_cells[1], 'Arial', 16)
            set_cell_font(row_cells[2], 'Arial', 16)

        blank_paragraph = doc.add_paragraph(f"Blank Votes: {contest.blank_votes}")
        over_paragraph = doc.add_paragraph(f"Over Votes: {contest.over_votes}")
        invalid_paragraph = doc.add_paragraph(f"Invalid Votes: {contest.invalid_votes}")

        for paragraph in [blank_paragraph, over_paragraph, invalid_paragraph]:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16)
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    return doc


def print_doc(file_path):
    # Get the default printer
    printer_name = win32print.GetDefaultPrinter()
    # Print the file
    print(f"Printing {file_path} on {printer_name} but not really")
    # win32api.ShellExecute(
    #     0,
    #     "print",
    #     file_path,
    #     f'/d:"{printer_name}"',
    #     ".",
    #     0
    # )


# Example usage
contests = [
    Contest("Contest 1", [Candidate("Candidate A", 300), Candidate("Candidate B", 150), Candidate("Candidate C", 50)],
            10, 5, 2),
    Contest("Contest 2", [Candidate("Candidate D", 500), Candidate("Candidate E", 300), Candidate("Candidate F", 100)],
            20, 3, 1)
]

# Generate OpenDoc file
doc = generate_opendoc(contests)
docx_file = "election_results.docx"
doc.save(docx_file)

# Optionally send the .docx file to the printer
print_doc(docx_file)
