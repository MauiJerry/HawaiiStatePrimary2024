from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK


# Assuming the structure of the Candidate and Contest classes

class Candidate:
    def __init__(self, name, votes):
        self.name = name
        self.votes = votes


class Contest:
    def __init__(self, name, candidates, blank_votes, over_votes, invalid_votes):
        self.name = name
        self.candidates = candidates
        self.blank_votes = blank_votes
        self.over_votes = over_votes
        self.invalid_votes = invalid_votes


def set_cell_font(cell, font_name, font_size, bold=False):
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def estimate_table_size(contest, font_size=18, padding=2):
    num_rows = len(contest.candidates) + 1  # Include header row
    estimated_row_height = font_size + padding
    total_height = num_rows * estimated_row_height

    estimated_col_widths = [
        max(len(candidate.name) for candidate in contest.candidates) * (font_size / 2),  # Candidate column
        len("Votes") * (font_size / 2),  # Votes column
        len("Percentage") * (font_size / 2)  # Percentage column
    ]

    return total_height

def generate_opendoc(contests):
    doc = Document()

    # Set page layout (optional, if you need to switch to landscape mode)
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Pt(595)  # A4 portrait width
    section.page_height = Pt(842)  # A4 portrait height
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    current_height = 0
    page_height_limit = (842 - 144)/2  # Page height minus top and bottom margins (in points)

    for contest in contests:
        estimated_height = estimate_table_size(contest)
        print(f"estimated height for contest {contest.name}= {estimated_height}")
        print(f"   {current_height}, {page_height_limit}")

        # Check if the contest fits on the current page
        if current_height + estimated_height > page_height_limit:
            doc.add_page_break()
            print("adding break")
            current_height = 0  # Reset the current height for the new page
        current_height += estimated_height  # Add the contest's height to the current page height
        print(f"page height now {current_height}")

        heading = doc.add_heading(contest.name, level=1)
        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Candidate'
        hdr_cells[1].text = 'Votes'
        hdr_cells[2].text = 'Percentage'

        set_cell_font(hdr_cells[0], 'Arial', 16, True)
        set_cell_font(hdr_cells[1], 'Arial', 16, True)
        set_cell_font(hdr_cells[2], 'Arial', 16, True)

        total_votes = sum(candidate.votes for candidate in contest.candidates)
        total_votes += contest.blank_votes + contest.over_votes + contest.invalid_votes

        sorted_candidates = sorted(contest.candidates, key=lambda x: x.votes, reverse=True)
        for candidate in sorted_candidates:
            row_cells = table.add_row().cells
            row_cells[0].text = candidate.name
            row_cells[1].text = str(candidate.votes)
            percentage = (candidate.votes / total_votes) * 100
            row_cells[2].text = f"{percentage:.2f}%"

            set_cell_font(row_cells[0], 'Arial', 16)
            set_cell_font(row_cells[1], 'Arial', 16)
            set_cell_font(row_cells[2], 'Arial', 16)

        blank_paragraph = doc.add_paragraph(f"Blank Votes: {contest.blank_votes}")
        over_paragraph = doc.add_paragraph(f"Over Votes: {contest.over_votes}")
        invalid_paragraph = doc.add_paragraph(f"Invalid Votes: {contest.invalid_votes}")

        for paragraph in [blank_paragraph, over_paragraph, invalid_paragraph]:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16)
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    return doc


# Example usage
contests = [
    Contest("Contest 1", [Candidate("Candidate A", 300), Candidate("Candidate B", 150), Candidate("Candidate C", 50)],
            10, 5, 2),
    Contest("Contest 2",
            [Candidate("Candidate D", 500), Candidate("Candidate E", 300), Candidate("Candidate F", 100),
                Candidate("Candidate A", 300), Candidate("Candidate B", 150), Candidate("Candidate C", 50) ],
            10, 5, 2),
    Contest("Contest 3",
            [Candidate("Candidate D", 500), Candidate("Candidate E", 300), Candidate("Candidate F", 100),
                Candidate("Candidate A", 300), Candidate("Candidate B", 150), Candidate("Candidate C", 50) ],
            15, 4, 3),
    Contest("Contest 4",
            [Candidate("Candidate G", 450), Candidate("Candidate H", 250),
             Candidate("Candidate I", 200),Candidate("Candidate G", 450),
             Candidate("Candidate H", 250), Candidate("Candidate I", 200)],
            15, 4, 3),
    Contest("Contest 5",
            [Candidate("Candidate G", 450), Candidate("Candidate H", 250), Candidate("Candidate I", 200)],
            15, 4, 3),
    Contest("Contest 6",
            [Candidate("Candidate G", 450), Candidate("Candidate H", 250), Candidate("Candidate I", 200)],
            15, 4, 3),
]

# Generate OpenDoc file
doc = generate_opendoc(contests)
docx_file = "election_results.docx"
doc.save(docx_file)

# Optionally, print the .docx file directly using the previously defined print_doc function
